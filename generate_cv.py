from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

# Fungsi untuk mengisi template
def generate_cv(data, template_path, output_path):
    # Buka template dokumen
    doc = Document(template_path)

    # Cari dan ganti placeholder di luar tabel
    for paragraph in doc.paragraphs:
        for key, value in data.items():
            # Mengganti {{Nama_Besar}} dengan format besar di bagian pertama
            if key == "Nama_Besar" and f"{{{{Nama_Besar}}}}" in paragraph.text:
                for run in paragraph.runs:
                    if f"{{{{Nama_Besar}}}}" in run.text:
                        run.text = run.text.replace(f"{{{{Nama_Besar}}}}", value)
                        run.font.size = Pt(18)
                        run.font.name = 'Arial'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
            # Ganti placeholder lainnya termasuk {{Nama}} yang tidak diformat besar
            elif f"{{{{{key}}}}}" in paragraph.text:
                paragraph.text = paragraph.text.replace(f"{{{{{key}}}}}", value)

    # Cari dan ganti placeholder di dalam tabel
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in data.items():
                    if f"{{{{{key}}}}}" in cell.text:
                        cell.text = cell.text.replace(f"{{{{{key}}}}}", value)

    # Simpan dokumen yang sudah diisi
    doc.save(output_path)
    print(f"Daftar Riwayat Hidup berhasil dibuat: {output_path}")

# Path template dan data dosen
template_path = "C:/Users/BTW-MODUL/Downloads/generate_cv.py/template_cv.docx"

# Data untuk 15 dosen
data_dosen_list = [
    {
        "Nama_Besar":"Dr. Alif Wibowo, S.T., M.T., Ph.D.",
        "Nama": "Dr. Alif Wibowo, S.T., M.T., Ph.D.",
        "NIP/NIK": "1234567890123456",
        "NIDN": "0123456789",
        "Tempat & Tanggal Lahir": "Surabaya, 15 Agustus 1980",
        "Jenis Kelamin": "Laki-laki",
        "Status Perkawinan": "Menikah",
        "Agama": "Islam",
        "Golongan / Pangkat": "IV/a – Pembina",
        "Jabatan Akademik": "Lektor Kepala",
        "Perguruan Tinggi": "Universitas Teknologi Indonesia",
        "Alamat": "Jl. Diponegoro No. 32, Denpasar, Bali",
        "Telp./Facs": "(0361) 123456",
        "Alamat Rumah": "Jl. Kenangan No. 17, Denpasar, Bali",
        "Telp./HP/Facs": "081234567890",
        "Alamat e-mail": "alif.wibowo@universitas.ac.id",
        # Pendidikan
        "Tahun_Lulus_1": "2006", "Program_Pendidikan_1": "Doktor", "Perguruan_Tinggi_1": "Universitas Teknologi Indonesia", "Jurusan_1": "Teknik Mesin",
        "Tahun_Lulus_2": "2002", "Program_Pendidikan_2": "Magister", "Perguruan_Tinggi_2": "Universitas Brawijaya", "Jurusan_2": "Teknik Industri",
        "Tahun_Lulus_3": "1998", "Program_Pendidikan_3": "Sarjana", "Perguruan_Tinggi_3": "Institut Teknologi Bandung", "Jurusan_3": "Teknik Sipil",
        "Tahun_Lulus_4": "1994", "Program_Pendidikan_4": "Diploma", "Perguruan_Tinggi_4": "Politeknik Negeri Jakarta", "Jurusan_4": "Teknik Mesin",
        # Pelatihan
        "Pelatihan_Tahun_1": "2019", "Jenis_Pelatihan_1": "Metode Penelitian", "Penyelenggara_1": "Lembaga Riset Nasional", "Sertifikat_1": "Sertifikat Penelitian", "Jangka_Waktu_1": "1 bulan",
        "Pelatihan_Tahun_2": "2021", "Jenis_Pelatihan_2": "Manajemen Proyek", "Penyelenggara_2": "Kominfo", "Sertifikat_2": "Sertifikat Manajemen Proyek", "Jangka_Waktu_2": "3 bulan",
        # Pengalaman Kerja
        "No_1": "1", "Pengalaman_Tahun_1": "2014-2024", "Jabatan_1": "Dosen Tetap", "Tempat_Bekerja_1": "Universitas Teknologi Indonesia",
        "No_2": "2", "Pengalaman_Tahun_2": "2009-2013", "Jabatan_2": "Asisten Dosen", "Tempat_Bekerja_2": "Institut Teknologi Bandung",
        "No_3": "3", "Pengalaman_Tahun_3": "2005-2008", "Jabatan_3": "Teknisi Laboratorium", "Tempat_Bekerja_3": "Politeknik Negeri Jakarta",
        # Penelitian
        "Penelitian_Tahun_1": "2021", "Judul_Penelitian_1": "Studi Material Ramah Lingkungan", "Peran_1": "Ketua", "Luaran_1": "Jurnal Nasional",
        "Penelitian_Tahun_2": "2018", "Judul_Penelitian_2": "Pengembangan Teknologi Berkelanjutan", "Peran_2": "Anggota", "Luaran_2": "Prosiding Konferensi",
        "Penelitian_Tahun_3": "2015", "Judul_Penelitian_3": "Teknologi Pengolahan Limbah", "Peran_3": "Ketua", "Luaran_3": "Jurnal Internasional",

        # Karya Ilmiah
        "Karya_Tahun_1": "2021", "Judul_Karya_1": "Buku Teknologi Mesin", "Penerbit_1": "Penerbit Ilmu Indonesia",
        "Karya_Tahun_2": "2018", "Judul_Karya_2": "Artikel Mesin Berbasis AI", "Penerbit_2": "Jurnal Teknologi",
        "Karya_Tahun_3": "2016", "Judul_Karya_3": "Makalah Mesin Otomotif", "Penerbit_3": "Jurnal Otomotif",
        # Makalah
       "Makalah_Tahun_1": "2020", "Judul_Makalah_1": "Pemanfaatan AI dalam Teknik Mesin", "Penyelenggara_Makalah_1": "Konferensi Nasional Teknik",
       "Makalah_Tahun_2": "2019", "Judul_Makalah_2": "Sistem Manufaktur Berbasis IoT", "Penyelenggara_Makalah_2": "Seminar Teknik Industri",
       "Makalah_Tahun_3": "2017", "Judul_Makalah_3": "Optimasi Produksi dengan Lean Manufacturing", "Penyelenggara_Makalah_3": "Lokakarya Lean Management",
        # Konferensi/Seminar
        "Konferensi_Tahun_1": "2022", "Judul_Kegiatan_1": "Seminar Nasional", "Penyelenggara_Konferensi_1": "Universitas Teknologi Indonesia", "Peran_Konferensi_1": "Peserta",
        "Konferensi_Tahun_2": "2021", "Judul_Kegiatan_2": "Simposium Teknologi AI", "Penyelenggara_Konferensi_2": "Institut Teknologi Bandung", "Peran_Konferensi_2": "Pembicara",
        "Konferensi_Tahun_3": "2020", "Judul_Kegiatan_3": "Lokakarya Robotika", "Penyelenggara_Konferensi_3": "Politeknik Negeri Jakarta", "Peran_Konferensi_3": "Panitia",

        # Pengabdian kepada Masyarakat
       "Kegiatan_Tahun_1": "2023", "Nama_Kegiatan_1": "Pelatihan Teknik Mesin untuk UMKM", "Tempat_Kegiatan_1": "Denpasar, Bali",
       "Kegiatan_Tahun_2": "2020", "Nama_Kegiatan_2": "Edukasi Kesehatan Lingkungan", "Tempat_Kegiatan_2": "Jakarta",
        # Penghargaan
       "Penghargaan_Tahun_1": "2020", "Bentuk_Penghargaan_1": "Dosen Berprestasi", "Pemberi_Penghargaan_1": "Universitas Teknologi Indonesia",
       "Penghargaan_Tahun_2": "2018", "Bentuk_Penghargaan_2": "Penghargaan Publikasi Internasional", "Pemberi_Penghargaan_2": "Kementerian Ristek",
       "Penghargaan_Tahun_3": "2016", "Bentuk_Penghargaan_3": "Inovator Teknologi Hijau", "Pemberi_Penghargaan_3": "Badan Pengembangan Teknologi",

        # Organisasi
        "Organisasi_Tahun_1": "2015-sekarang", "Nama_Organisasi_1": "Asosiasi Dosen Indonesia", "Jabatan_Organisasi_1": "Anggota",
        "Organisasi_Tahun_2": "2010-2015", "Nama_Organisasi_2": "Perhimpunan Teknik Mesin Indonesia", "Jabatan_Organisasi_2": "Sekretaris"
    },
    {
        "Nama_Besar":"Ir. Bela Harmanto, S.Kom., M.Eng",
        "Nama": "Ir. Bela Harmanto, S.Kom., M.Eng",
        "NIP/NIK": "9876543210987654",
        "NIDN": "9876543210",
        "Tempat & Tanggal Lahir": "Denpasar, 23 Juli 1975",
        "Jenis Kelamin": "Laki-laki",
        "Status Perkawinan": "Menikah",
        "Agama": "Hindu",
        "Golongan / Pangkat": "IV/b – Pembina Utama Muda",
        "Jabatan Akademik": "Guru Besar",
        "Perguruan Tinggi": "Universitas Hindu Indonesia",
        "Alamat": "Jl. Kusuma Bangsa No. 45, Sanur, Denpasar, Bali",
        "Telp./Facs": "(0361) 654321",
        "Alamat Rumah":"Jl. Merdeka No. 123, Kecamatan Sukajaya, Kota Bandung, Jawa Barat 40123",
        "Telp./HP/Facs": "081298765432",
        "Alamat e-mail": "bela.harmanto@institusi.ac.id",
        # Pendidikan
        "Tahun_Lulus_2": "2001", "Program_Pendidikan_2": "Magister", "Perguruan_Tinggi_2": "Universitas Udayana", "Jurusan_2": "Teknologi Informasi",
        "Tahun_Lulus_3": "1998", "Program_Pendidikan_3": "Sarjana", "Perguruan_Tinggi_3": "Institut Teknologi Bandung", "Jurusan_3": "Teknik Elektro",
        "Tahun_Lulus_4": "1994", "Program_Pendidikan_4": "Diploma", "Perguruan_Tinggi_4": "Politeknik Negeri Jakarta", "Jurusan_4": "Teknik Mesin",
        # Pelatihan
         "Pelatihan_Tahun_1": "2018", "Jenis_Pelatihan_1": "Hukum dan Keamanan Digital", "Penyelenggara_1": "Kemenkumham", "Sertifikat_1": "Sertifikat Hukum", "Jangka_Waktu_1": "2 minggu",
         "Pelatihan_Tahun_2": "2020", "Jenis_Pelatihan_2": "Manajemen Proyek Teknologi", "Penyelenggara_2": "Kominfo", "Sertifikat_2": "Sertifikat Manajemen Proyek", "Jangka_Waktu_2": "3 bulan",
        # Pengalaman Kerja
        "No_1": "1", "Pengalaman_Tahun_1": "2010-2022", "Jabatan_1": "Guru Besar", "Tempat_Bekerja_1": "Universitas Hindu Indonesia",
        "No_2": "2", "Pengalaman_Tahun_2": "2005-2010", "Jabatan_2": "Asisten Profesor", "Tempat_Bekerja_2": "Universitas Udayana",
        "No_3": "3", "Pengalaman_Tahun_3": "1998-2004", "Jabatan_3": "Teknisi Sistem Informasi", "Tempat_Bekerja_3": "Politeknik Negeri Jakarta",
        # Penelitian
        "Penelitian_Tahun_1": "2022", "Judul_Penelitian_1": "Kajian Hukum Adat Bali dalam Era Digital", "Peran_1": "Ketua", "Luaran_1": "Jurnal Internasional",
        "Penelitian_Tahun_2": "2020", "Judul_Penelitian_2": "Pengaruh Teknologi pada Sistem Hukum", "Peran_2": "Anggota", "Luaran_2": "Prosiding Konferensi",
        "Penelitian_Tahun_3": "2018", "Judul_Penelitian_3": "Etika Profesi di Era Teknologi", "Peran_3": "Ketua", "Luaran_3": "Jurnal Nasional",
        # Karya Ilmiah
         "Karya_Tahun_1": "2019", "Judul_Karya_1": "Buku Hukum Adat di Era Teknologi", "Penerbit_1": "Penerbit Saraswati",
         "Karya_Tahun_2": "2016", "Judul_Karya_2": "Artikel Teknologi dan Hukum", "Penerbit_2": "Jurnal Teknologi Hukum",
         "Karya_Tahun_3": "2013", "Judul_Karya_3": "Makalah Etika Profesi Hukum", "Penerbit_3": "Jurnal Hukum Nasional",
        # Makalah
        "Makalah_Tahun_1": "2020", "Judul_Makalah_1": "Peran Hukum dalam Teknologi Informasi", "Penyelenggara_Makalah_1": "Konferensi Hukum Nasional",
        "Makalah_Tahun_2": "2017", "Judul_Makalah_2": "Privasi dan Keamanan di Dunia Digital", "Penyelenggara_Makalah_2": "Seminar Nasional Teknologi",
        "Makalah_Tahun_3": "2014", "Judul_Makalah_3": "Implikasi Teknologi pada Etika Hukum", "Penyelenggara_Makalah_3": "Lokakarya Etika Hukum",
        # Konferensi/Seminar
        "Konferensi_Tahun_1": "2021", "Judul_Kegiatan_1": "Konferensi Hukum Internasional", "Penyelenggara_Konferensi_1": "Universitas Hindu Indonesia", "Peran_Konferensi_1": "Pembicara",
        "Konferensi_Tahun_2": "2018", "Judul_Kegiatan_2": "Simposium Keamanan Data", "Penyelenggara_Konferensi_2": "Universitas Indonesia", "Peran_Konferensi_2": "Peserta",
        "Konferensi_Tahun_3": "2015", "Judul_Kegiatan_3": "Lokakarya Teknologi Hukum", "Penyelenggara_Konferensi_3": "Politeknik Negeri Jakarta", "Peran_Konferensi_3": "Panitia",
        # Pengabdian kepada Masyarakat
         "Kegiatan_Tahun_1": "2020", "Nama_Kegiatan_1": "Sosialisasi Hukum dan Teknologi", "Tempat_Kegiatan_1": "Karangasem, Bali",
         "Kegiatan_Tahun_2": "2017", "Nama_Kegiatan_2": "Pelatihan Keamanan Data untuk Pelajar", "Tempat_Kegiatan_2": "Denpasar, Bali",
        # Penghargaan
        "Penghargaan_Tahun_1": "2018", "Bentuk_Penghargaan_1": "Penghargaan Guru Besar Terbaik", "Pemberi_Penghargaan_1": "Kementerian Pendidikan",
        "Penghargaan_Tahun_2": "2016", "Bentuk_Penghargaan_2": "Dosen Berprestasi", "Pemberi_Penghargaan_2": "Universitas Hindu Indonesia",
        "Penghargaan_Tahun_3": "2014", "Bentuk_Penghargaan_3": "Peneliti Hukum Teknologi Terbaik", "Pemberi_Penghargaan_3": "Badan Riset Nasional",
        # Organisasi
         "Organisasi_Tahun_1": "2010-sekarang", "Nama_Organisasi_1": "Ikatan Dosen Teknologi Hukum Indonesia", "Jabatan_Organisasi_1": "Sekretaris",
         "Organisasi_Tahun_2": "2005-2010", "Nama_Organisasi_2": "Perhimpunan Pakar Hukum Teknologi", "Jabatan_Organisasi_2": "Anggota"
    },
    {
        "Nama_Besar":"Dr. Fajar Kurniawan, S.T., M.Sc., Ph.D.",
        "Nama": "Dr. Fajar Kurniawan, S.T., M.Sc., Ph.D.",
        "NIP/NIK": "3456789012345678",
        "NIDN": "3456789012",
        "Tempat & Tanggal Lahir": "Jakarta, 17 April 1982",
        "Jenis Kelamin": "Laki-laki",
        "Status Perkawinan": "Menikah",
        "Agama": "Islam",
        "Golongan / Pangkat": "IV/a – Pembina",
        "Jabatan Akademik": "Lektor",
        "Perguruan Tinggi": "Institut Teknologi Bandung",
        "Alamat": "Jl. Melati No. 78, Jakarta Pusat",
        "Telp./Facs": "(021) 3456789",
        "Alamat Rumah": "Jl. Kenangan No. 34, Jakarta Pusat",
        "Telp./HP/Facs": "081123456789",
        "Alamat e-mail": "fajar.kurniawan@itb.ac.id",
        # Pendidikan
        "Tahun_Lulus_1": "2008", "Program_Pendidikan_1": "Doktor", "Perguruan_Tinggi_1": "Institut Teknologi Bandung", "Jurusan_1": "Teknik Elektro",
    "Tahun_Lulus_2": "2005", "Program_Pendidikan_2": "Magister", "Perguruan_Tinggi_2": "Universitas Indonesia", "Jurusan_2": "Teknik Elektro",
    "Tahun_Lulus_3": "2001", "Program_Pendidikan_3": "Sarjana", "Perguruan_Tinggi_3": "Universitas Gadjah Mada", "Jurusan_3": "Teknik Elektro",

        # Pelatihan
        "Pelatihan_Tahun_1": "2020", "Jenis_Pelatihan_1": "Teknologi Informasi", "Penyelenggara_1": "Kominfo", "Sertifikat_1": "Sertifikat IT", "Jangka_Waktu_1": "3 bulan",
        # Pengalaman Kerja
        "Pengalaman_Tahun_1": "2010-2020", "Jabatan_1": "Kepala Laboratorium", "Tempat_Bekerja_1": "Institut Teknologi Bandung",
        # Penelitian
        "Penelitian_Tahun_1": "2020", "Judul_Penelitian_1": "Inovasi Teknologi AI", "Peran_1": "Ketua", "Luaran_1": "Jurnal Nasional",
        # Karya Ilmiah
        "Karya_Tahun_1": "2019", "Judul_Karya_1": "Buku Teknologi AI", "Penerbit_1": "Penerbit ITB",
        # Konferensi/Seminar
        "Konferensi_Tahun_1": "2019", "Judul_Kegiatan_1": "Seminar AI", "Penyelenggara_Konferensi_1": "Kominfo", "Peran_Konferensi_1": "Pembicara",
        # Pengabdian kepada Masyarakat
        "Kegiatan_Tahun_1": "2018", "Nama_Kegiatan_1": "Workshop AI", "Tempat_Kegiatan_1": "Jakarta",
        # Penghargaan
        "Penghargaan_Tahun_1": "2017", "Bentuk_Penghargaan_1": "Dosen Berprestasi", "Pemberi_Penghargaan_1": "Institut Teknologi Bandung",
        # Organisasi
        "Organisasi_Tahun_1": "2015-sekarang", "Nama_Organisasi_1": "Asosiasi Ilmuwan IT", "Jabatan_Organisasi_1": "Anggota"
    },
    {
        "Nama_Besar":"Ir. Gita Prasetya, S.Kom., M.Eng.",
        "Nama": "Ir. Gita Prasetya, S.Kom., M.Eng.",
        "NIP/NIK": "4567890123456789",
        "NIDN": "4567890123",
        "Tempat & Tanggal Lahir": "Yogyakarta, 30 Juni 1985",
        "Jenis Kelamin": "Perempuan",
        "Status Perkawinan": "Belum Menikah",
        "Agama": "Islam",
        "Golongan / Pangkat": "III/c – Penata",
        "Jabatan Akademik": "Asisten Ahli",
        "Perguruan Tinggi": "Universitas Gadjah Mada",
        "Alamat": "Jl. Mawar No. 101, Yogyakarta",
        "Telp./Facs": "(0274) 987654",
        "Alamat Rumah": "Jl. Diponegoro No. 45, Yogyakarta",
        "Telp./HP/Facs": "081234567890",
        "Alamat e-mail": "gita.prasetya@ugm.ac.id",
        # Pendidikan
       "Tahun_Lulus_2": "2005", "Program_Pendidikan_2": "Magister", "Perguruan_Tinggi_2": "Universitas Indonesia", "Jurusan_2": "Teknik Elektro",
       "Tahun_Lulus_3": "2002", "Program_Pendidikan_3": "Sarjana", "Perguruan_Tinggi_3": "Institut Teknologi Sepuluh Nopember", "Jurusan_3": "Teknik Informatika",
       "Tahun_Lulus_4": "1998", "Program_Pendidikan_4": "Diploma", "Perguruan_Tinggi_4": "Politeknik Negeri Jakarta", "Jurusan_4": "Teknik Komputer",
        # Pelatihan
        "Pelatihan_Tahun_1": "2020", "Jenis_Pelatihan_1": "Teknologi Informasi", "Penyelenggara_1": "Kominfo", "Sertifikat_1": "Sertifikat IT", "Jangka_Waktu_1": "3 bulan",
        "Pelatihan_Tahun_2": "2018", "Jenis_Pelatihan_2": "Manajemen Proyek Teknologi", "Penyelenggara_2": "PMI Indonesia", "Sertifikat_2": "Sertifikat Manajemen Proyek", "Jangka_Waktu_2": "6 bulan",
        # Pengalaman Kerja
        "No_1": "1", "Pengalaman_Tahun_1": "2010-2020", "Jabatan_1": "Kepala Laboratorium", "Tempat_Bekerja_1": "Institut Teknologi Bandung",
        "No_2": "2", "Pengalaman_Tahun_2": "2006-2010", "Jabatan_2": "Koordinator Proyek TI", "Tempat_Bekerja_2": "Universitas Indonesia",
        "No_3": "3", "Pengalaman_Tahun_3": "2002-2006", "Jabatan_3": "Asisten Peneliti", "Tempat_Bekerja_3": "LIPI",
        # Penelitian
        "Penelitian_Tahun_1": "2020", "Judul_Penelitian_1": "Inovasi Teknologi AI", "Peran_1": "Ketua", "Luaran_1": "Jurnal Nasional",
        "Penelitian_Tahun_2": "2018", "Judul_Penelitian_2": "Integrasi IoT dalam Sistem Pendidikan", "Peran_2": "Anggota", "Luaran_2": "Prosiding Konferensi",
        "Penelitian_Tahun_3": "2015", "Judul_Penelitian_3": "Pengembangan Keamanan Data", "Peran_3": "Ketua", "Luaran_3": "Jurnal Internasional",
        # Karya Ilmiah
        "Karya_Tahun_1": "2019", "Judul_Karya_1": "Buku Teknologi AI", "Penerbit_1": "Penerbit ITB",
        "Karya_Tahun_2": "2017", "Judul_Karya_2": "Artikel Tentang IoT", "Penerbit_2": "Jurnal Teknologi Informasi",
        "Karya_Tahun_3": "2014", "Judul_Karya_3": "Makalah Sistem Embedded", "Penerbit_3": "Jurnal Embedded Systems",
        # Makalah
        "Makalah_Tahun_1": "2018", "Judul_Makalah_1": "AI dan Dampaknya pada Industri", "Penyelenggara_Makalah_1": "Seminar AI Nasional",
        "Makalah_Tahun_2": "2016", "Judul_Makalah_2": "IoT dalam Industri 4.0", "Penyelenggara_Makalah_2": "Simposium Teknologi ITB",
        "Makalah_Tahun_3": "2015", "Judul_Makalah_3": "Keamanan Data dalam Dunia Digital", "Penyelenggara_Makalah_3": "Konferensi Nasional TI",

        # Konferensi/Seminar
        "Konferensi_Tahun_1": "2019", "Judul_Kegiatan_1": "Seminar AI", "Penyelenggara_Konferensi_1": "Kominfo", "Peran_Konferensi_1": "Pembicara",
        "Konferensi_Tahun_2": "2017", "Judul_Kegiatan_2": "Simposium Keamanan Data", "Penyelenggara_Konferensi_2": "Universitas Indonesia", "Peran_Konferensi_2": "Panitia",
        "Konferensi_Tahun_3": "2016", "Judul_Kegiatan_3": "Lokakarya Teknologi IoT", "Penyelenggara_Konferensi_3": "Politeknik Negeri Jakarta", "Peran_Konferensi_3": "Peserta",
        # Pengabdian kepada Masyarakat
        "Kegiatan_Tahun_1": "2018", "Nama_Kegiatan_1": "Workshop AI untuk Siswa SMA", "Tempat_Kegiatan_1": "Jakarta",
        "Kegiatan_Tahun_2": "2016", "Nama_Kegiatan_2": "Sosialisasi Keamanan Internet", "Tempat_Kegiatan_2": "Bandung",
        # Penghargaan
        "Penghargaan_Tahun_1": "2017", "Bentuk_Penghargaan_1": "Dosen Berprestasi", "Pemberi_Penghargaan_1": "Institut Teknologi Bandung",
        "Penghargaan_Tahun_2": "2015", "Bentuk_Penghargaan_2": "Peneliti Terbaik", "Pemberi_Penghargaan_2": "Universitas Indonesia",
        "Penghargaan_Tahun_3": "2013", "Bentuk_Penghargaan_3": "Pakar Keamanan Data", "Pemberi_Penghargaan_3": "Kominfo",
        # Organisasi
        "Organisasi_Tahun_1": "2015-sekarang", "Nama_Organisasi_1": "Asosiasi Ilmuwan IT", "Jabatan_Organisasi_1": "Anggota",
        "Organisasi_Tahun_2": "2012-2015", "Nama_Organisasi_2": "Komunitas Keamanan Siber Indonesia", "Jabatan_Organisasi_2": "Sekretaris"
    },
    {
        "Nama_Besar":"Dr. Hanif Basuki, S.Si., M.Sc.",
        "Nama": "Dr. Hanif Basuki, S.Si., M.Sc.",
        "NIP/NIK": "5678901234567890",
        "NIDN": "5678901234",
        "Tempat & Tanggal Lahir": "Bandung, 25 November 1978",
        "Jenis Kelamin": "Laki-laki",
        "Status Perkawinan": "Menikah",
        "Agama": "Islam",
        "Golongan / Pangkat": "III/d – Penata Tingkat I",
        "Jabatan Akademik": "Lektor",
        "Perguruan Tinggi": "Universitas Padjadjaran",
        "Alamat": "Jl. Raya Bunga No. 88, Bandung",
        "Telp./Facs": "(022) 7654321",
        "Alamat Rumah": "Jl. Braga No. 5, Bandung",
        "Telp./HP/Facs": "081245678901",
        "Alamat e-mail": "hanif.basuki@unpad.ac.id",
        # Pendidikan
        "Tahun_Lulus_1": "2014", "Program_Pendidikan_1": "Doktor", "Perguruan_Tinggi_1": "Universitas Padjadjaran", "Jurusan_1": "Bioteknologi",
        "Tahun_Lulus_2": "2010", "Program_Pendidikan_2": "Magister", "Perguruan_Tinggi_2": "Universitas Padjadjaran", "Jurusan_2": "Biologi",
        "Tahun_Lulus_3": "2005", "Program_Pendidikan_3": "Sarjana", "Perguruan_Tinggi_3": "Institut Teknologi Bandung", "Jurusan_3": "Biologi",

        # Pelatihan
        "Pelatihan_Tahun_1": "2018", "Jenis_Pelatihan_1": "Analisis Data", "Penyelenggara_1": "Data Science Institute", "Sertifikat_1": "Sertifikat Data Science", "Jangka_Waktu_1": "2 bulan",
        "Pelatihan_Tahun_2": "2020", "Jenis_Pelatihan_2": "Bioteknologi Lingkungan", "Penyelenggara_2": "Pusat Bioteknologi Nasional", "Sertifikat_2": "Sertifikat Bioteknologi", "Jangka_Waktu_2": "3 bulan",
        # Pengalaman Kerja
        "No_1": "1", "Pengalaman_Tahun_1": "2015-2020", "Jabatan_1": "Peneliti Senior", "Tempat_Bekerja_1": "Universitas Padjadjaran",
        "No_2": "2", "Pengalaman_Tahun_2": "2010-2015", "Jabatan_2": "Asisten Peneliti", "Tempat_Bekerja_2": "Institut Teknologi Bandung",
        "No_3": "3", "Pengalaman_Tahun_3": "2008-2010", "Jabatan_3": "Koordinator Laboratorium", "Tempat_Bekerja_3": "LIPI",
        # Penelitian
       "Penelitian_Tahun_1": "2019", "Judul_Penelitian_1": "Pemanfaatan Enzim dalam Bioteknologi", "Peran_1": "Ketua", "Luaran_1": "Jurnal Internasional",
       "Penelitian_Tahun_2": "2016", "Judul_Penelitian_2": "Pengembangan Teknologi Fermentasi", "Peran_2": "Anggota", "Luaran_2": "Prosiding Nasional",
       "Penelitian_Tahun_3": "2013", "Judul_Penelitian_3": "Bioremediasi Limbah Cair", "Peran_3": "Ketua", "Luaran_3": "Jurnal Nasional",
        # Karya Ilmiah
        "Karya_Tahun_1": "2018", "Judul_Karya_1": "Makalah Enzim Bioteknologi", "Penerbit_1": "Journal of Biotechnology",
        "Karya_Tahun_2": "2015", "Judul_Karya_2": "Artikel Teknologi Pangan", "Penerbit_2": "Jurnal Teknologi Pangan",
        "Karya_Tahun_3": "2012", "Judul_Karya_3": "Buku tentang Genetika dan Lingkungan", "Penerbit_3": "Penerbit Ilmu Hayati",
        # Makalah
        "Makalah_Tahun_1": "2017", "Judul_Makalah_1": "Pemanfaatan Bioteknologi dalam Industri", "Penyelenggara_Makalah_1": "Konferensi Bioteknologi Nasional",
        "Makalah_Tahun_2": "2016", "Judul_Makalah_2": "Inovasi Pangan Berbasis Biologi Molekuler", "Penyelenggara_Makalah_2": "Seminar Teknologi Pangan",
        "Makalah_Tahun_3": "2014", "Judul_Makalah_3": "Fermentasi Enzimatik dalam Bioteknologi", "Penyelenggara_Makalah_3": "Lokakarya Bioteknologi",
        # Konferensi/Seminar
        "Konferensi_Tahun_1": "2018", "Judul_Kegiatan_1": "Bioteknologi Dunia", "Penyelenggara_Konferensi_1": "UNPAD", "Peran_Konferensi_1": "Pembicara",
        "Konferensi_Tahun_2": "2016", "Judul_Kegiatan_2": "Simposium Biologi Molekuler", "Penyelenggara_Konferensi_2": "Institut Teknologi Bandung", "Peran_Konferensi_2": "Panitia",
        "Konferensi_Tahun_3": "2014", "Judul_Kegiatan_3": "Lokakarya Genetika dan Lingkungan", "Penyelenggara_Konferensi_3": "LIPI", "Peran_Konferensi_3": "Peserta",
        # Pengabdian kepada Masyarakat
        "Kegiatan_Tahun_1": "2017", "Nama_Kegiatan_1": "Edukasi Bioteknologi", "Tempat_Kegiatan_1": "Bandung",
        "Kegiatan_Tahun_2": "2015", "Nama_Kegiatan_2": "Sosialisasi Teknologi Pangan Sehat", "Tempat_Kegiatan_2": "Yogyakarta",

        # Penghargaan
        "Penghargaan_Tahun_1": "2016", "Bentuk_Penghargaan_1": "Peneliti Terbaik", "Pemberi_Penghargaan_1": "Kementerian Ristek",
        "Penghargaan_Tahun_2": "2014", "Bentuk_Penghargaan_2": "Inovator Terbaik Bioteknologi", "Pemberi_Penghargaan_2": "UNPAD",
        "Penghargaan_Tahun_3": "2012", "Bentuk_Penghargaan_3": "Penghargaan Lingkungan Hidup", "Pemberi_Penghargaan_3": "Kementerian Lingkungan Hidup",
        # Organisasi
        "Organisasi_Tahun_1": "2010-sekarang", "Nama_Organisasi_1": "Asosiasi Bioteknologi", "Jabatan_Organisasi_1": "Anggota",
        "Organisasi_Tahun_2": "2008-2010", "Nama_Organisasi_2": "Himpunan Biologi Indonesia", "Jabatan_Organisasi_2": "Sekretaris"
    },
    {
        "Nama_Besar":"Ir. Indah Prihastuti, S.Kom., M.Kom.",
        "Nama": "Ir. Indah Prihastuti, S.Kom., M.Kom.",
        "NIP/NIK": "6789012345678901",
        "NIDN": "6789012345",
        "Tempat & Tanggal Lahir": "Solo, 10 September 1984",
        "Jenis Kelamin": "Perempuan",
        "Status Perkawinan": "Belum Menikah",
        "Agama": "Islam",
        "Golongan / Pangkat": "III/b – Penata Muda Tingkat I",
        "Jabatan Akademik": "Asisten Ahli",
        "Perguruan Tinggi": "Universitas Sebelas Maret",
        "Alamat": "Jl. Flamboyan No. 34, Solo",
        "Telp./Facs": "(0271) 234567",
        "Alamat Rumah": "Jl. Merdeka No. 56, Solo",
        "Telp./HP/Facs": "081234567892",
        "Alamat e-mail": "indah.prihastuti@uns.ac.id",
        # Pendidikan
       "Tahun_Lulus_1": "2016", "Program_Pendidikan_1": "Magister", "Perguruan_Tinggi_1": "Universitas Sebelas Maret", "Jurusan_1": "Informatika",
       "Tahun_Lulus_2": "2012", "Program_Pendidikan_2": "Sarjana", "Perguruan_Tinggi_2": "Universitas Gadjah Mada", "Jurusan_2": "Ilmu Komputer",

        # Pelatihan
        "Pelatihan_Tahun_1": "2019", "Jenis_Pelatihan_1": "Web Development", "Penyelenggara_1": "Code Academy", "Sertifikat_1": "Sertifikat Web Dev", "Jangka_Waktu_1": "3 bulan",
        "Pelatihan_Tahun_2": "2021", "Jenis_Pelatihan_2": "Cyber Security", "Penyelenggara_2": "Cyber Academy", "Sertifikat_2": "Sertifikat Cyber Security", "Jangka_Waktu_2": "4 bulan",
        # Pengalaman Kerja
        "No_1": "1", "Pengalaman_Tahun_1": "2018-2021", "Jabatan_1": "Developer", "Tempat_Bekerja_1": "PT Teknologi Nusantara",
        "No_2": "2", "Pengalaman_Tahun_2": "2016-2018", "Jabatan_2": "Junior Web Developer", "Tempat_Bekerja_2": "PT Digital Solution",
        "No_3": "3", "Pengalaman_Tahun_3": "2014-2016", "Jabatan_3": "Freelance Programmer", "Tempat_Bekerja_3": "Self-employed",

        # Penelitian
        "Penelitian_Tahun_1": "2020", "Judul_Penelitian_1": "Pengembangan Web Interaktif", "Peran_1": "Ketua", "Luaran_1": "Jurnal Nasional",
        "Penelitian_Tahun_2": "2018", "Judul_Penelitian_2": "Optimalisasi Sistem IoT", "Peran_2": "Anggota", "Luaran_2": "Prosiding Konferensi",
        "Penelitian_Tahun_3": "2017", "Judul_Penelitian_3": "Keamanan Data pada Sistem Web", "Peran_3": "Ketua", "Luaran_3": "Jurnal Internasional",
        # Karya Ilmiah
        "Karya_Tahun_1": "2021", "Judul_Karya_1": "Buku Panduan Web Dev", "Penerbit_1": "Penerbit Informatika",
        "Karya_Tahun_2": "2019", "Judul_Karya_2": "Artikel Tentang Keamanan Web", "Penerbit_2": "Jurnal Teknologi Informasi",
        "Karya_Tahun_3": "2018", "Judul_Karya_3": "Makalah Pengenalan HTML5", "Penerbit_3": "Jurnal Pengembangan Web",
        # Makalah
         "Makalah_Tahun_1": "2020", "Judul_Makalah_1": "Peningkatan User Experience dalam Desain Web", "Penyelenggara_Makalah_1": "Konferensi Nasional UX",
        "Makalah_Tahun_2": "2019", "Judul_Makalah_2": "Manajemen Data dalam Web Services", "Penyelenggara_Makalah_2": "Seminar Teknologi Informasi",
        "Makalah_Tahun_3": "2018", "Judul_Makalah_3": "Web Responsive Design", "Penyelenggara_Makalah_3": "Workshop Desain Web",

        # Konferensi/Seminar
        "Konferensi_Tahun_1": "2018", "Judul_Kegiatan_1": "Seminar Web Development", "Penyelenggara_Konferensi_1": "UNS", "Peran_Konferensi_1": "Pembicara",
        "Konferensi_Tahun_2": "2019", "Judul_Kegiatan_2": "Simposium Keamanan Siber", "Penyelenggara_Konferensi_2": "Universitas Indonesia", "Peran_Konferensi_2": "Peserta",
        "Konferensi_Tahun_3": "2020", "Judul_Kegiatan_3": "Lokakarya IoT dan Web Services", "Penyelenggara_Konferensi_3": "Kominfo", "Peran_Konferensi_3": "Pembicara",   "Konferensi_Tahun_3": "2020", "Judul_Kegiatan_3": "Lokakarya IoT dan Web Services", "Penyelenggara_Konferensi_3": "Kominfo", "Peran_Konferensi_3": "Pembicara",
        # Pengabdian kepada Masyarakat
        "Kegiatan_Tahun_1": "2019", "Nama_Kegiatan_1": "Pelatihan Web untuk UMKM", "Tempat_Kegiatan_1": "Solo",
        "Kegiatan_Tahun_2": "2020", "Nama_Kegiatan_2": "Sosialisasi Keamanan Data Online", "Tempat_Kegiatan_2": "Yogyakarta",
        # Penghargaan
        "Penghargaan_Tahun_1": "2021", "Bentuk_Penghargaan_1": "Web Developer Terbaik", "Pemberi_Penghargaan_1": "Universitas Sebelas Maret",
        "Penghargaan_Tahun_2": "2018", "Bentuk_Penghargaan_2": "Kontributor Teknologi Informasi", "Pemberi_Penghargaan_2": "Kominfo",
        "Penghargaan_Tahun_3": "2017", "Bentuk_Penghargaan_3": "Innovator of the Year", "Pemberi_Penghargaan_3": "Asosiasi Informatika Indonesia",
        # Organisasi
        "Organisasi_Tahun_1": "2015-sekarang", "Nama_Organisasi_1": "Asosiasi Informatika Indonesia", "Jabatan_Organisasi_1": "Anggota",
        "Organisasi_Tahun_2": "2018-2020", "Nama_Organisasi_2": "Himpunan Pengembang Web", "Jabatan_Organisasi_2": "Sekretaris"
    },
    {
         "Nama_Besar":"Dr. Johan Satria, S.T., M.Eng.",
        "Nama": "Dr. Johan Satria, S.T., M.Eng.",
        "NIP/NIK": "7890123456789012",
        "NIDN": "7890123456",
        "Tempat & Tanggal Lahir": "Medan, 2 Oktober 1981",
        "Jenis Kelamin": "Laki-laki",
        "Status Perkawinan": "Menikah",
        "Agama": "Islam",
        "Golongan / Pangkat": "IV/a – Pembina",
        "Jabatan Akademik": "Lektor Kepala",
        "Perguruan Tinggi": "Universitas Sumatera Utara",
        "Alamat": "Jl. Garuda No. 15, Medan",
        "Telp./Facs": "(061) 3456789",
        "Alamat Rumah": "Jl. Merpati No. 20, Medan",
        "Telp./HP/Facs": "081234567891",
        "Alamat e-mail": "johan.satria@usu.ac.id",
        # Pendidikan
        "Tahun_Lulus_1": "2015", "Program_Pendidikan_1": "Doktor", "Perguruan_Tinggi_1": "Universitas Indonesia", "Jurusan_1": "Teknik Sipil",
    "Tahun_Lulus_2": "2011", "Program_Pendidikan_2": "Magister", "Perguruan_Tinggi_2": "Institut Teknologi Bandung", "Jurusan_2": "Teknik Sipil",
    "Tahun_Lulus_3": "2007", "Program_Pendidikan_3": "Sarjana", "Perguruan_Tinggi_3": "Universitas Sumatera Utara", "Jurusan_3": "Teknik Sipil",

        # Pelatihan
        "Penelitian_Tahun_1": "2021", "Judul_Penelitian_1": "Pembangunan Berkelanjutan", "Peran_1": "Ketua", "Luaran_1": "Jurnal Internasional",
        "Penelitian_Tahun_2": "2019", "Judul_Penelitian_2": "Penggunaan Bahan Daur Ulang dalam Konstruksi", "Peran_2": "Anggota", "Luaran_2": "Prosiding Konferensi",    "Penelitian_Tahun_3": "2018", "Judul_Penelitian_3": "Optimasi Struktur Bangunan Ramah Lingkungan", "Peran_3": "Ketua", "Luaran_3": "Jurnal Nasional",
        # Pengalaman Kerja
        "No_1": "1", "Pengalaman_Tahun_1": "2012-2021", "Jabatan_1": "Kepala Proyek", "Tempat_Bekerja_1": "PT Wijaya Karya",
        "No_2": "2", "Pengalaman_Tahun_2": "2010-2012", "Jabatan_2": "Project Engineer", "Tempat_Bekerja_2": "PT Jaya Konstruksi",
        "No_3": "3", "Pengalaman_Tahun_3": "2007-2010", "Jabatan_3": "Supervisor Lapangan", "Tempat_Bekerja_3": "PT Adhi Karya",

        # Penelitian
       "Penelitian_Tahun_1": "2021", "Judul_Penelitian_1": "Pembangunan Berkelanjutan", "Peran_1": "Ketua", "Luaran_1": "Jurnal Internasional",
        "Penelitian_Tahun_2": "2019", "Judul_Penelitian_2": "Penggunaan Bahan Daur Ulang dalam Konstruksi", "Peran_2": "Anggota", "Luaran_2": "Prosiding Konferensi",
        "Penelitian_Tahun_3": "2018", "Judul_Penelitian_3": "Optimasi Struktur Bangunan Ramah Lingkungan", "Peran_3": "Ketua", "Luaran_3": "Jurnal Nasional",

        # Karya Ilmiah
        "Karya_Tahun_1": "2019", "Judul_Karya_1": "Buku Teknik Sipil", "Penerbit_1": "Penerbit Sipil Indonesia",
        "Karya_Tahun_2": "2017", "Judul_Karya_2": "Artikel Teknologi Beton", "Penerbit_2": "Jurnal Teknik",
        "Karya_Tahun_3": "2015", "Judul_Karya_3": "Makalah tentang Pondasi Bangunan", "Penerbit_3": "Jurnal Teknik Sipil",
         # Makalah
        "Makalah_Tahun_1": "2020", "Judul_Makalah_1": "Inovasi dalam Material Konstruksi", "Penyelenggara_Makalah_1": "Konferensi Nasional Teknik",
        "Makalah_Tahun_2": "2018", "Judul_Makalah_2": "Penggunaan Teknologi Hijau di Konstruksi", "Penyelenggara_Makalah_2": "Seminar Lingkungan",
        "Makalah_Tahun_3": "2016", "Judul_Makalah_3": "Analisis Struktur dan Stabilitas Bangunan", "Penyelenggara_Makalah_3": "Lokakarya Konstruksi",
        # Konferensi/Seminar
        "Konferensi_Tahun_1": "2018", "Judul_Kegiatan_1": "Seminar Teknik Sipil", "Penyelenggara_Konferensi_1": "ITS", "Peran_Konferensi_1": "Peserta",
        "Konferensi_Tahun_2": "2019", "Judul_Kegiatan_2": "Simposium Teknik Berkelanjutan", "Penyelenggara_Konferensi_2": "Institut Teknologi Bandung", "Peran_Konferensi_2": "Pembicara",
        "Konferensi_Tahun_3": "2020", "Judul_Kegiatan_3": "Lokakarya Teknik Beton", "Penyelenggara_Konferensi_3": "Universitas Sumatera Utara", "Peran_Konferensi_3": "Panitia",
        # Pengabdian kepada Masyarakat
        "Kegiatan_Tahun_1": "2020", "Nama_Kegiatan_1": "Pembangunan Desa", "Tempat_Kegiatan_1": "Binjai",
        "Kegiatan_Tahun_2": "2019", "Nama_Kegiatan_2": "Pelatihan Konstruksi untuk Pemuda", "Tempat_Kegiatan_2": "Medan",
        # Penghargaan
        "Penghargaan_Tahun_1": "2019", "Bentuk_Penghargaan_1": "Insinyur Berprestasi", "Pemberi_Penghargaan_1": "Kementerian PUPR",
        "Penghargaan_Tahun_2": "2018", "Bentuk_Penghargaan_2": "Penghargaan Inovasi Konstruksi", "Pemberi_Penghargaan_2": "Asosiasi Konstruksi Indonesia",
        "Penghargaan_Tahun_3": "2017", "Bentuk_Penghargaan_3": "Penerima Hibah Riset Terapan", "Pemberi_Penghargaan_3": "Kementerian Ristek",

        # Organisasi
        "Organisasi_Tahun_1": "2015-sekarang", "Nama_Organisasi_1": "Ikatan Insinyur Indonesia", "Jabatan_Organisasi_1": "Anggota",
        "Organisasi_Tahun_2": "2010-2015", "Nama_Organisasi_2": "Himpunan Ahli Konstruksi", "Jabatan_Organisasi_2": "Sekretaris"
    },
    {
        "Nama_Besar":"Dr. Karina Ningsih, S.T., M.T., Ph.D.",
        "Nama": "Dr. Karina Ningsih, S.T., M.T., Ph.D.",
        "NIP/NIK": "8901234567890123",
        "NIDN": "8901234567",
        "Tempat & Tanggal Lahir": "Makassar, 19 Desember 1979",
        "Jenis Kelamin": "Perempuan",
        "Status Perkawinan": "Menikah",
        "Agama": "Islam",
        "Golongan / Pangkat": "IV/b – Pembina Utama Muda",
        "Jabatan Akademik": "Guru Besar",
        "Perguruan Tinggi": "Universitas Hasanuddin",
        "Alamat": "Jl. Pattimura No. 42, Makassar",
        "Telp./Facs": "(0411) 2345678",
        "Alamat Rumah": "Jl. Sudirman No. 78, Makassar",
        "Telp./HP/Facs": "081987654321",
        "Alamat e-mail": "karina.ningsih@unhas.ac.id",
        # Pendidikan
        "Tahun_Lulus_1": "2006", "Program_Pendidikan_1": "Doktor", "Perguruan_Tinggi_1": "Universitas Gadjah Mada", "Jurusan_1": "Teknik Kimia",
        "Tahun_Lulus_2": "2002", "Program_Pendidikan_2": "Magister", "Perguruan_Tinggi_2": "Universitas Hasanuddin", "Jurusan_2": "Teknik Kimia",
        "Tahun_Lulus_3": "1998", "Program_Pendidikan_3": "Sarjana", "Perguruan_Tinggi_3": "Institut Teknologi Sepuluh Nopember", "Jurusan_3": "Teknik Kimia",

        # Pelatihan
       "Pelatihan_Tahun_1": "2017", "Jenis_Pelatihan_1": "Industri Kimia", "Penyelenggara_1": "BPPT", "Sertifikat_1": "Sertifikat Industri", "Jangka_Waktu_1": "1 tahun",
       "Pelatihan_Tahun_2": "2019", "Jenis_Pelatihan_2": "Keselamatan Proses Kimia", "Penyelenggara_2": "LIPI", "Sertifikat_2": "Sertifikat Keselamatan", "Jangka_Waktu_2": "6 bulan",
        # Pengalaman Kerja
         "No_1": "1", "Pengalaman_Tahun_1": "2010-2018", "Jabatan_1": "Konsultan Kimia", "Tempat_Bekerja_1": "PT Kimia Farma",
         "No_2": "2", "Pengalaman_Tahun_2": "2008-2010", "Jabatan_2": "Ahli Kimia", "Tempat_Bekerja_2": "Badan Pengkajian dan Penerapan Teknologi",
         "No_3": "3", "Pengalaman_Tahun_3": "2005-2008", "Jabatan_3": "Peneliti Junior", "Tempat_Bekerja_3": "Universitas Hasanuddin",
        # Penelitian
        "Penelitian_Tahun_1": "2019", "Judul_Penelitian_1": "Pengembangan Energi Terbarukan", "Peran_1": "Ketua", "Luaran_1": "Jurnal Nasional",
        "Penelitian_Tahun_2": "2017", "Judul_Penelitian_2": "Pemanfaatan Bahan Daur Ulang dalam Industri Kimia", "Peran_2": "Anggota", "Luaran_2": "Prosiding Internasional",
        "Penelitian_Tahun_3": "2015", "Judul_Penelitian_3": "Inovasi Bahan Ramah Lingkungan", "Peran_3": "Ketua", "Luaran_3": "Jurnal Internasional",

        # Karya Ilmiah
        "Karya_Tahun_1": "2018", "Judul_Karya_1": "Buku Kimia Terapan", "Penerbit_1": "Penerbit Kimia Indonesia",
        "Karya_Tahun_2": "2016", "Judul_Karya_2": "Artikel tentang Katalis dalam Reaksi Kimia", "Penerbit_2": "Jurnal Teknik Kimia",
        "Karya_Tahun_3": "2014", "Judul_Karya_3": "Makalah tentang Teknik Pengolahan Limbah", "Penerbit_3": "Jurnal Lingkungan",
        # Makalah
        "Makalah_Tahun_1": "2020", "Judul_Makalah_1": "Pengembangan Proses Kimia Ramah Lingkungan", "Penyelenggara_Makalah_1": "Konferensi Nasional Teknik Kimia",
        "Makalah_Tahun_2": "2018", "Judul_Makalah_2": "Efisiensi Proses Industri Kimia", "Penyelenggara_Makalah_2": "Seminar Kimia Terapan",
        "Makalah_Tahun_3": "2016", "Judul_Makalah_3": "Analisis Keberlanjutan Industri Kimia", "Penyelenggara_Makalah_3": "Lokakarya Kimia",
        # Konferensi/Seminar
        "Konferensi_Tahun_1": "2020", "Judul_Kegiatan_1": "Seminar Kimia", "Penyelenggara_Konferensi_1": "Universitas Indonesia", "Peran_Konferensi_1": "Pembicara",
        "Konferensi_Tahun_2": "2018", "Judul_Kegiatan_2": "Simposium Teknologi Kimia", "Penyelenggara_Konferensi_2": "BPPT", "Peran_Konferensi_2": "Peserta",
        "Konferensi_Tahun_3": "2017", "Judul_Kegiatan_3": "Workshop Proses Kimia Hijau", "Penyelenggara_Konferensi_3": "LIPI", "Peran_Konferensi_3": "Panitia",
        # Pengabdian kepada Masyarakat
        "Kegiatan_Tahun_1": "2021", "Nama_Kegiatan_1": "Sosialisasi Lingkungan", "Tempat_Kegiatan_1": "Makassar",
        "Kegiatan_Tahun_2": "2019", "Nama_Kegiatan_2": "Pelatihan Pengolahan Limbah untuk Masyarakat", "Tempat_Kegiatan_2": "Gowa",

        # Penghargaan
         "Penghargaan_Tahun_1": "2018", "Bentuk_Penghargaan_1": "Peneliti Terbaik", "Pemberi_Penghargaan_1": "Universitas Hasanuddin",
        "Penghargaan_Tahun_2": "2016", "Bentuk_Penghargaan_2": "Penghargaan Inovasi Lingkungan", "Pemberi_Penghargaan_2": "Kementerian Lingkungan Hidup",
        "Penghargaan_Tahun_3": "2015", "Bentuk_Penghargaan_3": "Penerima Hibah Penelitian", "Pemberi_Penghargaan_3": "Kementerian Ristek",

        # Organisasi
        "Organisasi_Tahun_1": "2010-sekarang", "Nama_Organisasi_1": "Perhimpunan Ahli Kimia", "Jabatan_Organisasi_1": "Anggota",
        "Organisasi_Tahun_2": "2005-2010", "Nama_Organisasi_2": "Asosiasi Kimiawan Indonesia", "Jabatan_Organisasi_2": "Sekretaris"
    }, {
        "Nama_Besar":"Ir. Leo Wardhana, S.Kom., M.Eng.",
        "Nama": "Ir. Leo Wardhana, S.Kom., M.Eng.",
        "NIP/NIK": "9012345678901234",
        "NIDN": "9012345678",
        "Tempat & Tanggal Lahir": "Semarang, 23 Maret 1983",
        "Jenis Kelamin": "Laki-laki",
        "Status Perkawinan": "Belum Menikah",
        "Agama": "Kristen",
        "Golongan / Pangkat": "III/a – Penata Muda",
        "Jabatan Akademik": "Asisten Ahli",
        "Perguruan Tinggi": "Universitas Diponegoro",
        "Alamat": "Jl. Kuningan No. 55, Semarang",
        "Telp./Facs": "(024) 3456789",
        "Alamat Rumah": "Jl. Merah Putih No. 99, Semarang",
        "Telp./HP/Facs": "081223456789",
        "Alamat e-mail": "leo.wardhana@undip.ac.id",
        # Pendidikan
         "Tahun_Lulus_1": "2014", "Program_Pendidikan_1": "Magister", "Perguruan_Tinggi_1": "Universitas Diponegoro", "Jurusan_1": "Informatika",
         "Tahun_Lulus_2": "2010", "Program_Pendidikan_2": "Sarjana", "Perguruan_Tinggi_2": "Universitas Gadjah Mada", "Jurusan_2": "Ilmu Komputer",

        # Pelatihan
       "Pelatihan_Tahun_1": "2016", "Jenis_Pelatihan_1": "Data Mining", "Penyelenggara_1": "Data Science Academy", "Sertifikat_1": "Sertifikat Data Mining", "Jangka_Waktu_1": "3 bulan",
       "Pelatihan_Tahun_2": "2019", "Jenis_Pelatihan_2": "Machine Learning", "Penyelenggara_2": "Kominfo", "Sertifikat_2": "Sertifikat Machine Learning", "Jangka_Waktu_2": "6 bulan",
        # Pengalaman Kerja
        "No_1": "1", "Pengalaman_Tahun_1": "2015-2021", "Jabatan_1": "Data Scientist", "Tempat_Bekerja_1": "PT. Data Nusantara",
        "No_2": "2", "Pengalaman_Tahun_2": "2012-2015", "Jabatan_2": "Analis Data", "Tempat_Bekerja_2": "Universitas Diponegoro",
        "No_3": "3", "Pengalaman_Tahun_3": "2010-2012", "Jabatan_3": "Asisten Peneliti", "Tempat_Bekerja_3": "Universitas Gadjah Mada",
        # Penelitian
        "Penelitian_Tahun_1": "2020", "Judul_Penelitian_1": "Algoritma Machine Learning", "Peran_1": "Ketua", "Luaran_1": "Jurnal Nasional",
        "Penelitian_Tahun_2": "2018", "Judul_Penelitian_2": "Pemodelan Data untuk UMKM", "Peran_2": "Anggota", "Luaran_2": "Prosiding Konferensi",
        "Penelitian_Tahun_3": "2016", "Judul_Penelitian_3": "Optimasi Data Mining", "Peran_3": "Ketua", "Luaran_3": "Jurnal Internasional",
        # Karya Ilmiah
        "Karya_Tahun_1": "2019", "Judul_Karya_1": "Makalah Machine Learning", "Penerbit_1": "Journal of Data Science",
        "Karya_Tahun_2": "2017", "Judul_Karya_2": "Artikel tentang Data Mining", "Penerbit_2": "Jurnal Teknologi",
        "Karya_Tahun_3": "2015", "Judul_Karya_3": "Buku Panduan Data Science", "Penerbit_3": "Penerbit Informatika",
        # Makalah
        "Makalah_Tahun_1": "2020", "Judul_Makalah_1": "Pemanfaatan AI dalam Data Science", "Penyelenggara_Makalah_1": "Konferensi Nasional Informatika",
        "Makalah_Tahun_2": "2018", "Judul_Makalah_2": "Pemodelan Data dalam Ekonomi", "Penyelenggara_Makalah_2": "Seminar Nasional Data",
        "Makalah_Tahun_3": "2016", "Judul_Makalah_3": "Pengembangan Algoritma Data Mining", "Penyelenggara_Makalah_3": "Lokakarya Teknologi Informasi",
        # Konferensi/Seminar
       "Konferensi_Tahun_1": "2018", "Judul_Kegiatan_1": "Workshop Data Science", "Penyelenggara_Konferensi_1": "Universitas Diponegoro", "Peran_Konferensi_1": "Peserta",
        "Konferensi_Tahun_2": "2019", "Judul_Kegiatan_2": "Konferensi Teknologi Informasi", "Penyelenggara_Konferensi_2": "Institut Teknologi Bandung", "Peran_Konferensi_2": "Pembicara",
        "Konferensi_Tahun_3": "2017", "Judul_Kegiatan_3": "Seminar Big Data", "Penyelenggara_Konferensi_3": "Data Science Indonesia", "Peran_Konferensi_3": "Panitia",

        # Pengabdian kepada Masyarakat
        "Kegiatan_Tahun_1": "2019", "Nama_Kegiatan_1": "Pelatihan Data Science untuk UMKM", "Tempat_Kegiatan_1": "Semarang",
        "Kegiatan_Tahun_2": "2017", "Nama_Kegiatan_2": "Workshop Pengolahan Data untuk Mahasiswa", "Tempat_Kegiatan_2": "Yogyakarta",
        # Penghargaan
        "Penghargaan_Tahun_1": "2018", "Bentuk_Penghargaan_1": "Data Scientist Terbaik", "Pemberi_Penghargaan_1": "Universitas Diponegoro",
        "Penghargaan_Tahun_2": "2016", "Bentuk_Penghargaan_2": "Penghargaan Inovasi Data Mining", "Pemberi_Penghargaan_2": "Kementerian Ristek",
        "Penghargaan_Tahun_3": "2015", "Bentuk_Penghargaan_3": "Penerima Hibah Riset Data Science", "Pemberi_Penghargaan_3": "Kominfo",
        # Organisasi
        "Organisasi_Tahun_1": "2017-sekarang", "Nama_Organisasi_1": "Asosiasi Data Scientist Indonesia", "Jabatan_Organisasi_1": "Anggota",
        "Organisasi_Tahun_2": "2013-2017", "Nama_Organisasi_2": "Komunitas Pengolahan Data Indonesia", "Jabatan_Organisasi_2": "Pengurus"
    },
    {
        "Nama_Besar":"Dr. Mita Anggraini, S.Si., M.Sc.",
        "Nama": "Dr. Mita Anggraini, S.Si., M.Sc.",
        "NIP/NIK": "0123456789012345",
        "NIDN": "0123456789",
        "Tempat & Tanggal Lahir": "Bogor, 5 Mei 1987",
        "Jenis Kelamin": "Perempuan",
        "Status Perkawinan": "Menikah",
        "Agama": "Islam",
        "Golongan / Pangkat": "III/b – Penata Muda",
        "Jabatan Akademik": "Lektor",
        "Perguruan Tinggi": "Institut Pertanian Bogor",
        "Alamat": "Jl. Pajajaran No. 12, Bogor",
        "Telp./Facs": "(0251) 1234567",
        "Alamat Rumah": "Jl. Merdeka No. 88, Bogor",
        "Telp./HP/Facs": "081987654320",
        "Alamat e-mail": "mita.anggraini@ipb.ac.id",
        # Pendidikan
        "Tahun_Lulus_1": "2020", "Program_Pendidikan_1": "Doktor", "Perguruan_Tinggi_1": "Institut Pertanian Bogor", "Jurusan_1": "Bioteknologi",
        "Tahun_Lulus_2": "2015", "Program_Pendidikan_2": "Magister", "Perguruan_Tinggi_2": "Institut Pertanian Bogor", "Jurusan_2": "Biologi",
        "Tahun_Lulus_3": "2011", "Program_Pendidikan_3": "Sarjana", "Perguruan_Tinggi_3": "Universitas Indonesia", "Jurusan_3": "Bioteknologi",

        # Pelatihan
       "Pelatihan_Tahun_1": "2017", "Jenis_Pelatihan_1": "Bioteknologi Lingkungan", "Penyelenggara_1": "Pusat Bioteknologi", "Sertifikat_1": "Sertifikat Lingkungan", "Jangka_Waktu_1": "4 bulan",
        "Pelatihan_Tahun_2": "2020", "Jenis_Pelatihan_2": "Pengelolaan Limbah Organik", "Penyelenggara_2": "Kementerian Lingkungan Hidup", "Sertifikat_2": "Sertifikat Pengelolaan Limbah", "Jangka_Waktu_2": "3 bulan",
        # Pengalaman Kerja
        "No_1": "1", "Pengalaman_Tahun_1": "2016-2021", "Jabatan_1": "Peneliti", "Tempat_Bekerja_1": "LIPI",
        "No_2": "2", "Pengalaman_Tahun_2": "2011-2015", "Jabatan_2": "Asisten Peneliti", "Tempat_Bekerja_2": "Institut Pertanian Bogor",
        "No_3": "3", "Pengalaman_Tahun_3": "2009-2011", "Jabatan_3": "Teknisi Laboratorium", "Tempat_Bekerja_3": "Universitas Indonesia",
        # Penelitian
        "Penelitian_Tahun_1": "2019", "Judul_Penelitian_1": "Penanganan Limbah Organik", "Peran_1": "Ketua", "Luaran_1": "Jurnal Nasional",
        "Penelitian_Tahun_2": "2017", "Judul_Penelitian_2": "Pemanfaatan Enzim dalam Biokonversi", "Peran_2": "Anggota", "Luaran_2": "Prosiding Konferensi",
        "Penelitian_Tahun_3": "2015", "Judul_Penelitian_3": "Pengembangan Bioproduk Lingkungan", "Peran_3": "Ketua", "Luaran_3": "Jurnal Internasional",
        # Karya Ilmiah
       "Karya_Tahun_1": "2018", "Judul_Karya_1": "Makalah Bioteknologi", "Penerbit_1": "Journal of Environmental Science",
        "Karya_Tahun_2": "2016", "Judul_Karya_2": "Artikel Teknologi Lingkungan", "Penerbit_2": "Jurnal Teknologi Indonesia",
        "Karya_Tahun_3": "2014", "Judul_Karya_3": "Buku Panduan Bioteknologi", "Penerbit_3": "Penerbit IPB Press",
        # Makalah
        "Makalah_Tahun_1": "2020", "Judul_Makalah_1": "Inovasi Teknologi Biokonversi", "Penyelenggara_Makalah_1": "Konferensi Teknologi Biologi",
        "Makalah_Tahun_2": "2018", "Judul_Makalah_2": "Pemanfaatan Limbah Industri", "Penyelenggara_Makalah_2": "Seminar Nasional Biologi",
        "Makalah_Tahun_3": "2016", "Judul_Makalah_3": "Teknologi Terapan Biologi", "Penyelenggara_Makalah_3": "Lokakarya Lingkungan Hidup",
        # Konferensi/Seminar
       "Konferensi_Tahun_1": "2017", "Judul_Kegiatan_1": "Konferensi Bioteknologi", "Penyelenggara_Konferensi_1": "IPB", "Peran_Konferensi_1": "Pembicara",
        "Konferensi_Tahun_2": "2019", "Judul_Kegiatan_2": "Simposium Lingkungan", "Penyelenggara_Konferensi_2": "Kementerian Lingkungan Hidup", "Peran_Konferensi_2": "Peserta",
        "Konferensi_Tahun_3": "2021", "Judul_Kegiatan_3": "Workshop Bioteknologi", "Penyelenggara_Konferensi_3": "LIPI", "Peran_Konferensi_3": "Panitia",

        # Pengabdian kepada Masyarakat
        "Kegiatan_Tahun_1": "2018", "Nama_Kegiatan_1": "Edukasi Lingkungan", "Tempat_Kegiatan_1": "Bogor",
        "Kegiatan_Tahun_2": "2020", "Nama_Kegiatan_2": "Penyuluhan Kesehatan Lingkungan", "Tempat_Kegiatan_2": "Jakarta",
        # Penghargaan
        "Penghargaan_Tahun_1": "2019", "Bentuk_Penghargaan_1": "Peneliti Terbaik", "Pemberi_Penghargaan_1": "Institut Pertanian Bogor",
        "Penghargaan_Tahun_2": "2017", "Bentuk_Penghargaan_2": "Penghargaan Publikasi Lingkungan", "Pemberi_Penghargaan_2": "LIPI",
        "Penghargaan_Tahun_3": "2016", "Bentuk_Penghargaan_3": "Inovator Teknologi Biokonversi", "Pemberi_Penghargaan_3": "Kementerian Ristek",

        # Organisasi
       "Organisasi_Tahun_1": "2016-sekarang", "Nama_Organisasi_1": "Perkumpulan Bioteknologi", "Jabatan_Organisasi_1": "Anggota",
        "Organisasi_Tahun_2": "2013-2016", "Nama_Organisasi_2": "Asosiasi Peneliti Lingkungan", "Jabatan_Organisasi_2": "Sekretaris"
    },{
         "Nama_Besar":"Ir. Nanda Kusuma, S.Kom., M.Kom.",
        "Nama": "Ir. Nanda Kusuma, S.Kom., M.Kom.",
        "NIP/NIK": "2345678901234567",
        "NIDN": "2345678901",
        "Tempat & Tanggal Lahir": "Palembang, 12 Juni 1985",
        "Jenis Kelamin": "Laki-laki",
        "Status Perkawinan": "Menikah",
        "Agama": "Islam",
        "Golongan / Pangkat": "III/c – Penata",
        "Jabatan Akademik": "Lektor",
        "Perguruan Tinggi": "Universitas Sriwijaya",
        "Alamat": "Jl. Sudirman No. 65, Palembang",
        "Telp./Facs": "(0711) 987654",
        "Alamat Rumah": "Jl. Mawar No. 23, Palembang",
        "Telp./HP/Facs": "081298765432",
        "Alamat e-mail": "nanda.kusuma@unsri.ac.id",
        # Pendidikan
        "Tahun_Lulus_1": "2015", "Program_Pendidikan_1": "Magister", "Perguruan_Tinggi_1": "Universitas Sriwijaya", "Jurusan_1": "Sistem Informasi",
        "Tahun_Lulus_2": "2010", "Program_Pendidikan_2": "Sarjana", "Perguruan_Tinggi_2": "Universitas Indonesia", "Jurusan_2": "Ilmu Komputer",

        # Pelatihan
        "Pelatihan_Tahun_1": "2018", "Jenis_Pelatihan_1": "Cyber Security", "Penyelenggara_1": "Cyber Academy", "Sertifikat_1": "Sertifikat Cyber Security", "Jangka_Waktu_1": "3 bulan",
        "Pelatihan_Tahun_2": "2019", "Jenis_Pelatihan_2": "Data Protection", "Penyelenggara_2": "Kominfo", "Sertifikat_2": "Sertifikat Data Protection", "Jangka_Waktu_2": "2 bulan",
        # Pengalaman Kerja
        "No_1": "1", "Pengalaman_Tahun_1": "2016-2021", "Jabatan_1": "Spesialis Keamanan Siber", "Tempat_Bekerja_1": "Universitas Sriwijaya",
        "No_2": "2", "Pengalaman_Tahun_2": "2010-2015", "Jabatan_2": "Analis Data", "Tempat_Bekerja_2": "PT Telekomunikasi Indonesia",
        "No_3": "3", "Pengalaman_Tahun_3": "2008-2010", "Jabatan_3": "Teknisi Jaringan", "Tempat_Bekerja_3": "PT Jasa Telekomunikasi",
        # Penelitian
        "Penelitian_Tahun_1": "2020", "Judul_Penelitian_1": "Keamanan Data dalam Dunia Digital", "Peran_1": "Ketua", "Luaran_1": "Jurnal Nasional",
        "Penelitian_Tahun_2": "2018", "Judul_Penelitian_2": "Implementasi Enkripsi Data", "Peran_2": "Anggota", "Luaran_2": "Prosiding Konferensi",
        "Penelitian_Tahun_3": "2016", "Judul_Penelitian_3": "Keamanan Jaringan dalam Big Data", "Peran_3": "Ketua", "Luaran_3": "Jurnal Internasional",
        # Karya Ilmiah
        "Karya_Tahun_1": "2019", "Judul_Karya_1": "Makalah tentang Keamanan Data", "Penerbit_1": "Journal of Cyber Security",
        "Karya_Tahun_2": "2017", "Judul_Karya_2": "Artikel Keamanan Siber", "Penerbit_2": "Jurnal Sistem Informasi Indonesia",
        "Karya_Tahun_3": "2015", "Judul_Karya_3": "Buku Panduan Keamanan Jaringan", "Penerbit_3": "Penerbit Unsri Press",
        # Makalah
        "Makalah_Tahun_1": "2019", "Judul_Makalah_1": "Keamanan Data Pribadi di Era Digital", "Penyelenggara_Makalah_1": "Konferensi Cyber Security",
        "Makalah_Tahun_2": "2018", "Judul_Makalah_2": "Enkripsi Data untuk UMKM", "Penyelenggara_Makalah_2": "Seminar Nasional Keamanan Siber",
        "Makalah_Tahun_3": "2016", "Judul_Makalah_3": "Teknologi Perlindungan Jaringan", "Penyelenggara_Makalah_3": "Lokakarya Data Science",

        # Konferensi/Seminar
        "Konferensi_Tahun_1": "2017", "Judul_Kegiatan_1": "Seminar Keamanan Data", "Penyelenggara_Konferensi_1": "UNSRI", "Peran_Konferensi_1": "Peserta",
        "Konferensi_Tahun_2": "2019", "Judul_Kegiatan_2": "Simposium Teknologi Digital", "Penyelenggara_Konferensi_2": "Universitas Indonesia", "Peran_Konferensi_2": "Pembicara",
        "Konferensi_Tahun_3": "2021", "Judul_Kegiatan_3": "Workshop Sistem Informasi", "Penyelenggara_Konferensi_3": "Universitas Gadjah Mada", "Peran_Konferensi_3": "Panitia",

        # Pengabdian kepada Masyarakat
        "Kegiatan_Tahun_1": "2018", "Nama_Kegiatan_1": "Workshop Cyber Security", "Tempat_Kegiatan_1": "Palembang",
        "Kegiatan_Tahun_2": "2020", "Nama_Kegiatan_2": "Pelatihan Keamanan Data untuk UMKM", "Tempat_Kegiatan_2": "Medan",
        # Penghargaan
         "Penghargaan_Tahun_1": "2019", "Bentuk_Penghargaan_1": "Pakar Keamanan Siber", "Pemberi_Penghargaan_1": "Universitas Sriwijaya",
        "Penghargaan_Tahun_2": "2018", "Bentuk_Penghargaan_2": "Inovasi Keamanan Data", "Pemberi_Penghargaan_2": "Kementerian Komunikasi dan Informatika",
        "Penghargaan_Tahun_3": "2016", "Bentuk_Penghargaan_3": "Penerapan Teknologi Jaringan", "Pemberi_Penghargaan_3": "Asosiasi Teknologi Indonesia",
        # Organisasi
        "Organisasi_Tahun_1": "2016-sekarang", "Nama_Organisasi_1": "Asosiasi Keamanan Data", "Jabatan_Organisasi_1": "Anggota",
        "Organisasi_Tahun_2": "2014-2016", "Nama_Organisasi_2": "Perhimpunan Profesional Sistem Informasi", "Jabatan_Organisasi_2": "Sekretaris"
    },
    {
        "Nama_Besar":"Dr. Oscar Wijaya, S.T., M.Eng.",
        "Nama": "Dr. Oscar Wijaya, S.T., M.Eng.",
        "NIP/NIK": "3456789012345678",
        "NIDN": "3456789012",
        "Tempat & Tanggal Lahir": "Balikpapan, 14 September 1980",
        "Jenis Kelamin": "Laki-laki",
        "Status Perkawinan": "Menikah",
        "Agama": "Kristen",
        "Golongan / Pangkat": "IV/a – Pembina",
        "Jabatan Akademik": "Lektor Kepala",
        "Perguruan Tinggi": "Universitas Mulawarman",
        "Alamat": "Jl. Antasari No. 45, Balikpapan",
        "Telp./Facs": "(0542) 345678",
        "Alamat Rumah": "Jl. Martadinata No. 89, Balikpapan",
        "Telp./HP/Facs": "081345678912",
        "Alamat e-mail": "oscar.wijaya@unmul.ac.id",
        # Pendidikan
         "Tahun_Lulus_1": "2012", "Program_Pendidikan_1": "Doktor", "Perguruan_Tinggi_1": "Universitas Indonesia", "Jurusan_1": "Teknik Sipil",
    "Tahun_Lulus_2": "2008", "Program_Pendidikan_2": "Magister", "Perguruan_Tinggi_2": "Institut Teknologi Sepuluh Nopember", "Jurusan_2": "Teknik Sipil",
    "Tahun_Lulus_3": "2004", "Program_Pendidikan_3": "Sarjana", "Perguruan_Tinggi_3": "Universitas Mulawarman", "Jurusan_3": "Teknik Sipil",
        # Pelatihan
        "Pelatihan_Tahun_1": "2016", "Jenis_Pelatihan_1": "Konstruksi Ramah Lingkungan", "Penyelenggara_1": "LPJK", "Sertifikat_1": "Sertifikat Konstruksi Hijau", "Jangka_Waktu_1": "6 bulan",
        "Pelatihan_Tahun_2": "2018", "Jenis_Pelatihan_2": "Manajemen Proyek Konstruksi", "Penyelenggara_2": "Kominfo", "Sertifikat_2": "Sertifikat Manajemen Proyek", "Jangka_Waktu_2": "4 bulan",
        # Pengalaman Kerja
        "No_1": "1", "Pengalaman_Tahun_1": "2011-2020", "Jabatan_1": "Konsultan Sipil", "Tempat_Bekerja_1": "PT Wijaya Karya",
        "No_2": "2", "Pengalaman_Tahun_2": "2008-2011", "Jabatan_2": "Pengawas Proyek", "Tempat_Bekerja_2": "PT Adhi Karya",
        "No_3": "3", "Pengalaman_Tahun_3": "2004-2008", "Jabatan_3": "Staf Teknik", "Tempat_Bekerja_3": "PT Pembangunan Jaya",

        # Penelitian
        "Penelitian_Tahun_1": "2019", "Judul_Penelitian_1": "Penggunaan Bahan Daur Ulang dalam Konstruksi", "Peran_1": "Ketua", "Luaran_1": "Jurnal Internasional",
        "Penelitian_Tahun_2": "2016", "Judul_Penelitian_2": "Teknologi Konstruksi Hijau", "Peran_2": "Anggota", "Luaran_2": "Prosiding Konferensi",
        "Penelitian_Tahun_3": "2014", "Judul_Penelitian_3": "Efisiensi Energi pada Bangunan", "Peran_3": "Ketua", "Luaran_3": "Jurnal Nasional",
        # Karya Ilmiah
        "Karya_Tahun_1": "2018", "Judul_Karya_1": "Makalah Konstruksi Hijau", "Penerbit_1": "Journal of Sustainable Engineering",
        "Karya_Tahun_2": "2016", "Judul_Karya_2": "Buku Teknologi Konstruksi", "Penerbit_2": "Penerbit Teknik Indonesia",
        "Karya_Tahun_3": "2015", "Judul_Karya_3": "Artikel tentang Bangunan Hijau", "Penerbit_3": "Jurnal Teknik Lingkungan",
 # Makalah
    "Makalah_Tahun_1": "2019", "Judul_Makalah_1": "Efisiensi Energi pada Bangunan Hijau", "Penyelenggara_Makalah_1": "Seminar Konstruksi Hijau",
    "Makalah_Tahun_2": "2017", "Judul_Makalah_2": "Daur Ulang Bahan Konstruksi", "Penyelenggara_Makalah_2": "Konferensi Teknik Ramah Lingkungan",
    "Makalah_Tahun_3": "2015", "Judul_Makalah_3": "Inovasi Material Bangunan", "Penyelenggara_Makalah_3": "Lokakarya Teknologi Konstruksi",
        # Konferensi/Seminar
        "Konferensi_Tahun_1": "2019", "Judul_Kegiatan_1": "Seminar Lingkungan Konstruksi", "Penyelenggara_Konferensi_1": "UNMUL", "Peran_Konferensi_1": "Pembicara",
    "Konferensi_Tahun_2": "2018", "Judul_Kegiatan_2": "Simposium Bangunan Hijau", "Penyelenggara_Konferensi_2": "ITS", "Peran_Konferensi_2": "Panitia",
    "Konferensi_Tahun_3": "2016", "Judul_Kegiatan_3": "Workshop Konstruksi Berkelanjutan", "Penyelenggara_Konferensi_3": "LPJK", "Peran_Konferensi_3": "Peserta",

        # Pengabdian kepada Masyarakat
        "Kegiatan_Tahun_1": "2017", "Nama_Kegiatan_1": "Workshop Konstruksi Hijau", "Tempat_Kegiatan_1": "Balikpapan",
    "Kegiatan_Tahun_2": "2015", "Nama_Kegiatan_2": "Penyuluhan Lingkungan", "Tempat_Kegiatan_2": "Samarinda",

        # Penghargaan
       "Penghargaan_Tahun_1": "2018", "Bentuk_Penghargaan_1": "Pakar Konstruksi Hijau", "Pemberi_Penghargaan_1": "Universitas Mulawarman",
    "Penghargaan_Tahun_2": "2016", "Bentuk_Penghargaan_2": "Inovator Material Bangunan", "Pemberi_Penghargaan_2": "Kementerian PUPR",
    "Penghargaan_Tahun_3": "2015", "Bentuk_Penghargaan_3": "Kontributor Lingkungan Konstruksi", "Pemberi_Penghargaan_3": "Asosiasi Teknik Sipil Indonesia",

        # Organisasi
       "Organisasi_Tahun_1": "2012-sekarang", "Nama_Organisasi_1": "Ikatan Insinyur Sipil Indonesia", "Jabatan_Organisasi_1": "Ketua",
    "Organisasi_Tahun_2": "2010-2012", "Nama_Organisasi_2": "Perhimpunan Teknik Konstruksi Hijau", "Jabatan_Organisasi_2": "Sekretaris"
    },
    {
        "Nama_Besar":"Dr. Chandra Purnomo, S.Si., M.Sc.",
        "Nama": "Dr. Chandra Purnomo, S.Si., M.Sc.",
        "NIP/NIK": "5678901234567890",
        "NIDN": "5678901234",
        "Tempat & Tanggal Lahir": "Manado, 9 Juli 1975",
        "Jenis Kelamin": "Laki-Laki",
        "Status Perkawinan": "Menikah",
        "Agama": "Kristen",
        "Golongan / Pangkat": "IV/a – Pembina",
        "Jabatan Akademik": "Lektor Kepala",
        "Perguruan Tinggi": "Universitas Sam Ratulangi",
        "Alamat": "Jl. Wolter Monginsidi No. 23, Manado",
        "Telp./Facs": "(0431) 654321",
        "Alamat Rumah":"Jl. Pahlawan No. 45, Kelurahan Sukamaju, Kecamatan Menteng, Kota Jakarta Pusat, DKI Jakarta 10310",
        "Telp./HP/Facs": "081234567891",
        "Alamat e-mail": "chandra.purnomo@unsrat.ac.id",
        # Pendidikan
         "Tahun_Lulus_1": "2015", "Program_Pendidikan_1": "Doktor", "Perguruan_Tinggi_1": "Universitas Indonesia", "Jurusan_1": "Teknik Mesin",
    "Tahun_Lulus_2": "2010", "Program_Pendidikan_2": "Magister", "Perguruan_Tinggi_2": "Universitas Gadjah Mada", "Jurusan_2": "Teknik Mesin",
    "Tahun_Lulus_3": "2005", "Program_Pendidikan_3": "Sarjana", "Perguruan_Tinggi_3": "Universitas Sam Ratulangi", "Jurusan_3": "Teknik Mesin",
    
        # Pelatihan
        "Pelatihan_Tahun_1": "2019", "Jenis_Pelatihan_1": "Pemeliharaan Mesin", "Penyelenggara_1": "Lembaga Teknik Indonesia", "Sertifikat_1": "Sertifikat Pemeliharaan", "Jangka_Waktu_1": "3 bulan",
    "Pelatihan_Tahun_2": "2018", "Jenis_Pelatihan_2": "Manajemen Pemeliharaan", "Penyelenggara_2": "Kominfo", "Sertifikat_2": "Sertifikat Manajemen", "Jangka_Waktu_2": "2 bulan",
        # Pengalaman Kerja
        "No_1": "1", "Pengalaman_Tahun_1": "2015-2020", "Jabatan_1": "Pengawas Teknik", "Tempat_Bekerja_1": "PT. Astra Indonesia",
    "No_2": "2", "Pengalaman_Tahun_2": "2010-2015", "Jabatan_2": "Teknisi Mesin", "Tempat_Bekerja_2": "PT. Garuda Mesin",
    "No_3": "3", "Pengalaman_Tahun_3": "2005-2010", "Jabatan_3": "Staf Pemeliharaan", "Tempat_Bekerja_3": "PT. Cipta Mesin",
        # Penelitian
         "Penelitian_Tahun_1": "2020", "Judul_Penelitian_1": "Optimasi Pemeliharaan Mesin", "Peran_1": "Ketua", "Luaran_1": "Jurnal Nasional",
    "Penelitian_Tahun_2": "2017", "Judul_Penelitian_2": "Pemeliharaan Terjadwal untuk Industri", "Peran_2": "Anggota", "Luaran_2": "Prosiding Konferensi",
    "Penelitian_Tahun_3": "2015", "Judul_Penelitian_3": "Teknik Pemeliharaan Mesin Ramah Lingkungan", "Peran_3": "Ketua", "Luaran_3": "Jurnal Internasional",
        # Karya Ilmiah
        "Karya_Tahun_1": "2018", "Judul_Karya_1": "Makalah tentang Pemeliharaan Mesin", "Penerbit_1": "Journal of Mechanical Engineering",
    "Karya_Tahun_2": "2016", "Judul_Karya_2": "Buku Teknik Pemeliharaan Mesin", "Penerbit_2": "Penerbit Teknik Indonesia",
    "Karya_Tahun_3": "2014", "Judul_Karya_3": "Artikel Efisiensi Mesin", "Penerbit_3": "Jurnal Teknik Industri",
    # Makalah
    "Makalah_Tahun_1": "2019", "Judul_Makalah_1": "Pemeliharaan Mesin Modern", "Penyelenggara_Makalah_1": "Seminar Teknik Mesin Nasional",
    "Makalah_Tahun_2": "2018", "Judul_Makalah_2": "Optimalisasi Pemeliharaan pada Industri Manufaktur", "Penyelenggara_Makalah_2": "Konferensi Teknologi Manufaktur",
    "Makalah_Tahun_3": "2016", "Judul_Makalah_3": "Manajemen Pemeliharaan pada Industri Kecil", "Penyelenggara_Makalah_3": "Lokakarya Industri Kecil",

        # Konferensi/Seminar
        "Konferensi_Tahun_1": "2019", "Judul_Kegiatan_1": "Konferensi Nasional Teknik Mesin", "Penyelenggara_Konferensi_1": "UGM", "Peran_Konferensi_1": "Pembicara",
    "Konferensi_Tahun_2": "2017", "Judul_Kegiatan_2": "Simposium Manajemen Mesin", "Penyelenggara_Konferensi_2": "ITS", "Peran_Konferensi_2": "Panitia",
    "Konferensi_Tahun_3": "2016", "Judul_Kegiatan_3": "Workshop Pemeliharaan Mesin Industri", "Penyelenggara_Konferensi_3": "Lembaga Teknik Indonesia", "Peran_Konferensi_3": "Peserta",

        # Pengabdian kepada Masyarakat
        "Kegiatan_Tahun_1": "2017", "Nama_Kegiatan_1": "Pelatihan Teknik Mesin untuk Pemula", "Tempat_Kegiatan_1": "Manado",
    "Kegiatan_Tahun_2": "2015", "Nama_Kegiatan_2": "Sosialisasi Pemeliharaan Mesin", "Tempat_Kegiatan_2": "Tomohon",
        # Penghargaan
        "Penghargaan_Tahun_1": "2019", "Bentuk_Penghargaan_1": "Penghargaan Dosen Terbaik", "Pemberi_Penghargaan_1": "Universitas Sam Ratulangi",
    "Penghargaan_Tahun_2": "2017", "Bentuk_Penghargaan_2": "Kontributor Teknologi Pemeliharaan", "Pemberi_Penghargaan_2": "Asosiasi Teknik Mesin Indonesia",
    "Penghargaan_Tahun_3": "2015", "Bentuk_Penghargaan_3": "Pakar Pemeliharaan Industri", "Pemberi_Penghargaan_3": "Lembaga Teknik Indonesia",
        # Organisasi
        "Organisasi_Tahun_1": "2016-sekarang", "Nama_Organisasi_1": "Perhimpunan Teknik Mesin Indonesia", "Jabatan_Organisasi_1": "Anggota",
    "Organisasi_Tahun_2": "2010-2016", "Nama_Organisasi_2": "Asosiasi Pemeliharaan Mesin", "Jabatan_Organisasi_2": "Sekretaris"
    },
    {
         "Nama_Besar":"Dr. Eri Utami, S.T., M.Eng.",
        "Nama": "Dr. Eri Utami, S.T., M.Eng.",
        "NIP/NIK": "6789012345678901",
        "NIDN": "6789012345",
        "Tempat & Tanggal Lahir": "Surabaya, 20 November 1980",
        "Jenis Kelamin": "Laki-laki",
        "Status Perkawinan": "Menikah",
        "Agama": "Islam",
        "Golongan / Pangkat": "III/d – Penata Tingkat I",
        "Jabatan Akademik": "Lektor",
        "Perguruan Tinggi": "Institut Teknologi Sepuluh Nopember",
        "Alamat": "Jl. A. Yani No. 88, Surabaya",
        "Telp./Facs": "(031) 9876543",
        "Alamat Rumah": (
    "Jl. Anggrek No. 56, Kelurahan Pondok Indah, Kecamatan Kebayoran Lama, "
    "Kota Jakarta Selatan, DKI Jakarta 12310"
),
        "Telp./HP/Facs": "081245678902",
        "Alamat e-mail": "satria.pratama@its.ac.id",
        # Pendidikan
        "Tahun_Lulus_1": "2018", "Program_Pendidikan_1": "Doktor", "Perguruan_Tinggi_1": "Institut Teknologi Sepuluh Nopember", "Jurusan_1": "Informatika",
    "Tahun_Lulus_2": "2012", "Program_Pendidikan_2": "Magister", "Perguruan_Tinggi_2": "Institut Teknologi Sepuluh Nopember", "Jurusan_2": "Informatika",
    "Tahun_Lulus_3": "2008", "Program_Pendidikan_3": "Sarjana", "Perguruan_Tinggi_3": "Universitas Airlangga", "Jurusan_3": "Ilmu Komputer",
        # Pelatihan
         "Pelatihan_Tahun_1": "2018", "Jenis_Pelatihan_1": "Machine Learning", "Penyelenggara_1": "AI Indonesia", "Sertifikat_1": "Sertifikat Machine Learning", "Jangka_Waktu_1": "6 bulan",
    "Pelatihan_Tahun_2": "2016", "Jenis_Pelatihan_2": "Data Analysis", "Penyelenggara_2": "Big Data Academy", "Sertifikat_2": "Sertifikat Data Analysis", "Jangka_Waktu_2": "4 bulan",
        # Pengalaman Kerja
        "No_1": "1", "Pengalaman_Tahun_1": "2013-2021", "Jabatan_1": "Data Scientist", "Tempat_Bekerja_1": "Tokopedia",
    "No_2": "2", "Pengalaman_Tahun_2": "2011-2013", "Jabatan_2": "Analyst Programmer", "Tempat_Bekerja_2": "PT. Teknologi Nusantara",
    "No_3": "3", "Pengalaman_Tahun_3": "2008-2011", "Jabatan_3": "Junior Developer", "Tempat_Bekerja_3": "Bank Mandiri",
        # Penelitian
        "Penelitian_Tahun_1": "2020", "Judul_Penelitian_1": "Implementasi AI dalam Bisnis", "Peran_1": "Ketua", "Luaran_1": "Jurnal Internasional",
    "Penelitian_Tahun_2": "2018", "Judul_Penelitian_2": "Analisis Data Besar untuk E-commerce", "Peran_2": "Anggota", "Luaran_2": "Prosiding Konferensi",
    "Penelitian_Tahun_3": "2015", "Judul_Penelitian_3": "Optimasi Algoritma dalam Machine Learning", "Peran_3": "Ketua", "Luaran_3": "Jurnal Nasional",
        # Karya Ilmiah
        "Karya_Tahun_1": "2018", "Judul_Karya_1": "Makalah Machine Learning", "Penerbit_1": "Journal of Data Science",
    "Karya_Tahun_2": "2017", "Judul_Karya_2": "Buku Pengantar Data Science", "Penerbit_2": "Penerbit ITS",
    "Karya_Tahun_3": "2016", "Judul_Karya_3": "Artikel Implementasi AI", "Penerbit_3": "Jurnal Teknologi",
     # Makalah
    "Makalah_Tahun_1": "2019", "Judul_Makalah_1": "Penggunaan AI untuk Bisnis Retail", "Penyelenggara_Makalah_1": "Konferensi AI Nasional",
    "Makalah_Tahun_2": "2018", "Judul_Makalah_2": "Implementasi Machine Learning dalam Logistik", "Penyelenggara_Makalah_2": "Seminar Data Science Indonesia",
    "Makalah_Tahun_3": "2017", "Judul_Makalah_3": "Prediksi Bisnis Menggunakan Algoritma AI", "Penyelenggara_Makalah_3": "Lokakarya Teknologi Bisnis",
        # Konferensi/Seminar
        "Konferensi_Tahun_1": "2017", "Judul_Kegiatan_1": "Workshop AI", "Penyelenggara_Konferensi_1": "ITS", "Peran_Konferensi_1": "Peserta",
    "Konferensi_Tahun_2": "2015", "Judul_Kegiatan_2": "Simposium Teknologi Data", "Penyelenggara_Konferensi_2": "Universitas Airlangga", "Peran_Konferensi_2": "Panitia",
    "Konferensi_Tahun_3": "2014", "Judul_Kegiatan_3": "Seminar Nasional Teknologi Informasi", "Penyelenggara_Konferensi_3": "Kementerian Kominfo", "Peran_Konferensi_3": "Pembicara",
        # Pengabdian kepada Masyarakat
       "Kegiatan_Tahun_1": "2019", "Nama_Kegiatan_1": "Pelatihan AI untuk Startup", "Tempat_Kegiatan_1": "Surabaya",
    "Kegiatan_Tahun_2": "2016", "Nama_Kegiatan_2": "Workshop Data Science untuk Masyarakat", "Tempat_Kegiatan_2": "Malang",

        # Penghargaan
        "Penghargaan_Tahun_1": "2019", "Bentuk_Penghargaan_1": "Data Scientist Berprestasi", "Pemberi_Penghargaan_1": "ITS",
    "Penghargaan_Tahun_2": "2017", "Bentuk_Penghargaan_2": "Inovator Machine Learning", "Pemberi_Penghargaan_2": "AI Indonesia",
    "Penghargaan_Tahun_3": "2015", "Bentuk_Penghargaan_3": "Pengembang Teknologi Terbaik", "Pemberi_Penghargaan_3": "Universitas Airlangga",
        # Organisasi
        "Organisasi_Tahun_1": "2016-sekarang", "Nama_Organisasi_1": "Perkumpulan Data Science Indonesia", "Jabatan_Organisasi_1": "Anggota",
    "Organisasi_Tahun_2": "2014-2016", "Nama_Organisasi_2": "Asosiasi Teknologi Informatika", "Jabatan_Organisasi_2": "Sekretaris"

    },
    {
        "Nama_Besar":"Ir. Dewi Anastasya, S.Kom., M.Kom.",
        "Nama": "Ir. Dewi Anastasya, S.Kom., M.Kom.",
        "NIP/NIK": "7890123456789012",
        "NIDN": "7890123456",
        "Tempat & Tanggal Lahir": "Yogyakarta, 15 April 1982",
        "Jenis Kelamin": "Perempuani",
        "Status Perkawinan": "Menikah",
        "Agama": "Islam",
        "Golongan / Pangkat": "IV/a – Pembina",
        "Jabatan Akademik": "Lektor Kepala",
        "Perguruan Tinggi": "Universitas Gadjah Mada",
        "Alamat": "Jl. Kaliurang No. 101, Yogyakarta",
        "Telp./Facs": "(0274) 123456",
        "Alamat Rumah": (
    "Jl. Kenanga No. 89, Kelurahan Taman Sari, "
    "Kecamatan Taman, Kota Surabaya, Jawa Timur 60234"
),
        "Telp./HP/Facs": "081234567891",
        "Alamat e-mail": "wahyu.setiawan@ugm.ac.id",
        # Pendidikan
        "Tahun_Lulus_1": "2012", "Program_Pendidikan_1": "Magister", "Perguruan_Tinggi_1": "Universitas Gadjah Mada", "Jurusan_1": "Teknik Elektro",
        "Tahun_Lulus_2": "2008", "Program_Pendidikan_2": "Sarjana", "Perguruan_Tinggi_2": "Universitas Indonesia", "Jurusan_2": "Teknik Elektro",
        # Pelatihan
        "Pelatihan_Tahun_1": "2016", "Jenis_Pelatihan_1": "Internet of Things", "Penyelenggara_1": "Kominfo", "Sertifikat_1": "Sertifikat IoT", "Jangka_Waktu_1": "3 bulan",
        "Pelatihan_Tahun_2": "2015", "Jenis_Pelatihan_2": "Jaringan Komunikasi", "Penyelenggara_2": "Telkom Indonesia", "Sertifikat_2": "Sertifikat Jaringan", "Jangka_Waktu_2": "2 bulan",
        # Pengalaman Kerja
        "No_1": "1", "Pengalaman_Tahun_1": "2013-2021", "Jabatan_1": "Engineer IoT", "Tempat_Bekerja_1": "PT Telkom",
        "No_2": "2", "Pengalaman_Tahun_2": "2011-2013", "Jabatan_2": "Network Engineer", "Tempat_Bekerja_2": "PT Indosat",
        "No_3": "3", "Pengalaman_Tahun_3": "2008-2011", "Jabatan_3": "Asisten Teknisi", "Tempat_Bekerja_3": "PT Huawei Indonesia",

        # Penelitian
         "Penelitian_Tahun_1": "2020", "Judul_Penelitian_1": "Pengembangan Jaringan 5G", "Peran_1": "Ketua", "Luaran_1": "Jurnal Internasional",
    "Penelitian_Tahun_2": "2018", "Judul_Penelitian_2": "Keamanan IoT untuk Infrastruktur Kritis", "Peran_2": "Anggota", "Luaran_2": "Prosiding Konferensi",
    "Penelitian_Tahun_3": "2015", "Judul_Penelitian_3": "Optimasi Bandwidth dalam Sistem IoT", "Peran_3": "Ketua", "Luaran_3": "Jurnal Nasional",

        # Karya Ilmiah
        "Karya_Tahun_1": "2019", "Judul_Karya_1": "Makalah Teknologi IoT", "Penerbit_1": "Journal of Electrical Engineering",
    "Karya_Tahun_2": "2017", "Judul_Karya_2": "Buku Panduan IoT", "Penerbit_2": "Penerbit UGM",
    "Karya_Tahun_3": "2016", "Judul_Karya_3": "Artikel tentang 5G dan IoT", "Penerbit_3": "Jurnal Telekomunikasi",
     # Makalah
    "Makalah_Tahun_1": "2018", "Judul_Makalah_1": "Manfaat IoT dalam Industri Manufaktur", "Penyelenggara_Makalah_1": "Konferensi IoT Nasional",
    "Makalah_Tahun_2": "2016", "Judul_Makalah_2": "Teknologi Sensor untuk Lingkungan", "Penyelenggara_Makalah_2": "Seminar Teknologi Lingkungan",
    "Makalah_Tahun_3": "2015", "Judul_Makalah_3": "Jaringan IoT di Wilayah Terpencil", "Penyelenggara_Makalah_3": "Lokakarya Telekomunikasi",
        # Konferensi/Seminar
        "Konferensi_Tahun_1": "2018", "Judul_Kegiatan_1": "Konferensi Nasional IoT", "Penyelenggara_Konferensi_1": "UGM", "Peran_Konferensi_1": "Pembicara",
    "Konferensi_Tahun_2": "2017", "Judul_Kegiatan_2": "Simposium Teknologi Jaringan", "Penyelenggara_Konferensi_2": "Universitas Indonesia", "Peran_Konferensi_2": "Panitia",
    "Konferensi_Tahun_3": "2015", "Judul_Kegiatan_3": "Seminar Nasional Telekomunikasi", "Penyelenggara_Konferensi_3": "Kementerian Kominfo", "Peran_Konferensi_3": "Pembicara",
        # Pengabdian kepada Masyarakat
       "Kegiatan_Tahun_1": "2017", "Nama_Kegiatan_1": "Workshop IoT untuk Pelajar", "Tempat_Kegiatan_1": "Yogyakarta",
    "Kegiatan_Tahun_2": "2015", "Nama_Kegiatan_2": "Pelatihan Teknologi Jaringan untuk UMKM", "Tempat_Kegiatan_2": "Bandung",

        # Penghargaan
        "Penghargaan_Tahun_1": "2018", "Bentuk_Penghargaan_1": "Pakar Teknologi IoT", "Pemberi_Penghargaan_1": "Kementerian Kominfo",
    "Penghargaan_Tahun_2": "2016", "Bentuk_Penghargaan_2": "Inovator Teknologi Jaringan", "Pemberi_Penghargaan_2": "PT Telkom",
    "Penghargaan_Tahun_3": "2014", "Bentuk_Penghargaan_3": "Dosen Berprestasi", "Pemberi_Penghargaan_3": "Universitas Indonesia",
        # Organisasi
        "Organisasi_Tahun_1": "2015-sekarang", "Nama_Organisasi_1": "Perkumpulan IoT Indonesia", "Jabatan_Organisasi_1": "Anggota",
    "Organisasi_Tahun_2": "2013-2015", "Nama_Organisasi_2": "Ikatan Insinyur Elektro", "Jabatan_Organisasi_2": "Sekretaris"
    }
]

# Generate CV untuk setiap dosen dalam data_dosen_list
for i, data_dosen in enumerate(data_dosen_list):
    output_path = f"C:/Users/BTW-MODUL/Downloads/generate_cv.py/Daftar_Riwayat_Hidup_{data_dosen['Nama'].replace(' ', '_').replace('.', '')}.docx"
    generate_cv(data_dosen, template_path, output_path)